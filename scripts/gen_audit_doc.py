#!/usr/bin/env python3
"""
Generate security audit report in DOCX and PDF formats
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import subprocess
import os

# Audit data
audit_data = {
    "summary": "Based on checking the RustSec advisory database for key dependencies, all vulnerabilities are patched in current versions.",
    "license": "MIT License - Copyright (c) 2026 Mr. Yotsawarit Pudpong",
    "repository": "https://github.com/Yotsawarit/bnn-code.git",
    "packages": [
        {"package": "anyhow", "version": "1.0.104", "patched_in": ">= 1.0.103", "advisory": "RUSTSEC-2026-0190", "status": "[PATCHED]", "details": "Unsoundness in Error::downcast_mut() - memory corruption"},
        {"package": "chrono", "version": "0.4.45", "patched_in": ">= 0.4.20", "advisory": "RUSTSEC-2020-0159", "status": "[PATCHED]", "details": "Potential segfault in localtime_r invocations"},
        {"package": "tokio", "version": "1.53.1", "patched_in": ">= 1.44.2", "advisory": "RUSTSEC-2025-0023", "status": "[PATCHED]", "details": "Broadcast channel calls clone in parallel without Sync"},
        {"package": "regex", "version": "1.13.1", "patched_in": ">= 1.5.5", "advisory": "RUSTSEC-2022-0013", "status": "[PATCHED]", "details": "DoS via large repetitions on empty sub-expressions"},
        {"package": "rusqlite", "version": "0.31.0", "patched_in": ">= 0.26.2", "advisory": "RUSTSEC-2021-0128", "status": "[PATCHED]", "details": "Incorrect lifetime bounds on closures - use-after-free"},
        {"package": "tracing", "version": "0.1.44", "patched_in": ">= 0.1.40", "advisory": "RUSTSEC-2023-0078", "status": "[PATCHED]", "details": "Stack use-after-free in Instrumented::into_inner"},
        {"package": "rustls", "version": "0.23.40", "patched_in": ">= 0.23.18", "advisory": "RUSTSEC-2024-0399", "status": "[PATCHED]", "details": "Network-reachable panic in Acceptor::accept"},
        {"package": "hyper", "version": "1.11.0", "patched_in": ">= 0.14.10", "advisory": "RUSTSEC-2021-0078", "status": "[PATCHED]", "details": "Lenient Content-Length parsing - request smuggling"},
        {"package": "dirs", "version": "5.0.1", "patched_in": "N/A", "advisory": "RUSTSEC-2020-0053", "status": "[UNMAINTAINED]", "details": "Crate unmaintained - migrate to dirs-next"},
    ],
    "recommendation": "All critical vulnerabilities are patched. Only informational issue: consider migrating from 'dirs' to 'dirs-next' crate.",
    "scanned_by": "Manual RustSec advisory database check (cargo-audit installation pending)"
}

def create_docx():
    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Security Audit Report - bnn-code', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    doc.add_paragraph('Project: bnn-code v0.1.3')
    doc.add_paragraph('Scan Date: 2026-08-18')
    doc.add_paragraph('Scanner: RustSec Advisory Database (manual check)')
    doc.add_paragraph(f'Repository: {audit_data["repository"]}')
    doc.add_paragraph('')
    
    # Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(audit_data["summary"])
    doc.add_paragraph('')
    
    # Table
    doc.add_heading('Dependency Vulnerability Status', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    headers = ['Package', 'Current Version', 'Patched In', 'Advisory ID', 'Status', 'Details']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    # Data rows
    for pkg in audit_data["packages"]:
        row = table.add_row()
        row.cells[0].text = pkg["package"]
        row.cells[1].text = pkg["version"]
        row.cells[2].text = pkg["patched_in"]
        row.cells[3].text = pkg["advisory"]
        row.cells[4].text = pkg["status"]
        row.cells[5].text = pkg["details"]
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    # Set column widths
    widths = [Inches(0.8), Inches(0.9), Inches(0.9), Inches(1.1), Inches(0.7), Inches(2.5)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width
    
    doc.add_paragraph('')
    
    # Recommendation
    doc.add_heading('Recommendation', level=1)
    doc.add_paragraph(audit_data["recommendation"])
    doc.add_paragraph('')
    
    # Note
    doc.add_heading('Note', level=1)
    doc.add_paragraph(audit_data["scanned_by"])
    doc.add_paragraph('Run "cargo install cargo-audit && cargo audit" for automated scanning in CI/CD.')
    doc.add_paragraph('')
    
    # License
    doc.add_heading('License', level=1)
    doc.add_paragraph(audit_data["license"])
    doc.add_paragraph('Permission is hereby granted, free of charge, to any person obtaining a copy')
    doc.add_paragraph('of this software and associated documentation files (the "Software"), to deal')
    doc.add_paragraph('in the Software without restriction, including without limitation the rights')
    doc.add_paragraph('to use, copy, modify, merge, publish, distribute, sublicense, and/or sell')
    doc.add_paragraph('copies of the Software, and to permit persons to whom the Software is')
    doc.add_paragraph('furnished to do so, subject to the following conditions:')
    doc.add_paragraph('')
    doc.add_paragraph('The above copyright notice and this permission notice shall be included in all')
    doc.add_paragraph('copies or substantial portions of the Software.')
    doc.add_paragraph('')
    doc.add_paragraph('THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR')
    doc.add_paragraph('IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,')
    doc.add_paragraph('FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE')
    doc.add_paragraph('AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER')
    doc.add_paragraph('LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,')
    doc.add_paragraph('OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE')
    doc.add_paragraph('SOFTWARE.')
    
    # Save
    output_path = '/home/bnn-core/bnn-code/SECURITY_AUDIT.docx'
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")
    return output_path

def create_pdf_simple():
    """Create simple PDF using fpdf2 without complex tables"""
    from fpdf import FPDF
    
    class PDF(FPDF):
        def header(self):
            self.set_font('Helvetica', 'B', 16)
            self.cell(0, 10, 'Security Audit Report - bnn-code', new_x="LMARGIN", new_y="NEXT", align='C')
            self.set_font('Helvetica', '', 10)
            self.cell(0, 6, 'Project: bnn-code v0.1.3', new_x="LMARGIN", new_y="NEXT", align='C')
            self.cell(0, 6, 'Scan Date: 2026-08-18', new_x="LMARGIN", new_y="NEXT", align='C')
            self.cell(0, 6, f'Repository: {audit_data["repository"]}', new_x="LMARGIN", new_y="NEXT", align='C')
            self.ln(5)
        
        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')
    
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(15, 15, 15)
    
    # Summary
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Executive Summary', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 6, audit_data["summary"])
    pdf.ln(5)
    
    # Table as simple list
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Dependency Vulnerability Status', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    
    for pkg in audit_data["packages"]:
        pdf.set_font('Helvetica', 'B', 11)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 7, f'{pkg["package"]} v{pkg["version"]}')
        pdf.set_font('Helvetica', '', 10)
        pdf.set_x(pdf.l_margin + 5)
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 5, 6, f'Advisory: {pkg["advisory"]}')
        pdf.set_x(pdf.l_margin + 5)
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 5, 6, f'Status: {pkg["status"]}')
        pdf.set_x(pdf.l_margin + 5)
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 5, 6, f'Patched in: {pkg["patched_in"]}')
        pdf.set_x(pdf.l_margin + 5)
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 5, 6, f'Details: {pkg["details"]}')
        pdf.ln(4)
        
        if pdf.get_y() > 260:
            pdf.add_page()
    
    pdf.ln(5)
    
    # Recommendation
    if pdf.get_y() > 200:
        pdf.add_page()
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Recommendation', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 6, audit_data["recommendation"])
    pdf.ln(5)
    
    # Note
    if pdf.get_y() > 200:
        pdf.add_page()
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Note', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font('Helvetica', '', 11)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 6, audit_data["scanned_by"])
    if pdf.get_y() > 260:
        pdf.add_page()
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 6, 'Run "cargo install cargo-audit && cargo audit" for automated scanning in CI/CD.')
    
    # License
    if pdf.get_y() > 200:
        pdf.add_page()
    pdf.ln(5)
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'License', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 6, audit_data["license"])
    pdf.set_font('Helvetica', '', 10)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'Permission is hereby granted, free of charge, to any person obtaining a copy')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'of this software and associated documentation files (the "Software"), to deal')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'in the Software without restriction, including without limitation the rights')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'to use, copy, modify, merge, publish, distribute, sublicense, and/or sell')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'copies of the Software, and to permit persons to whom the Software is')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'furnished to do so, subject to the following conditions:')
    pdf.ln(2)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'The above copyright notice and this permission notice shall be included in all')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'copies or substantial portions of the Software.')
    pdf.ln(2)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE')
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 5, 'SOFTWARE.')
    
    pdf_path = '/home/bnn-core/bnn-code/SECURITY_AUDIT.pdf'
    pdf.output(pdf_path)
    print(f"PDF saved to: {pdf_path}")
    return pdf_path

if __name__ == '__main__':
    # Generate DOCX
    docx_path = create_docx()
    
    # Generate PDF
    pdf_path = create_pdf_simple()
    
    print("\nDone! Files created:")
    print(f"  - {docx_path}")
    print(f"  - {pdf_path}")